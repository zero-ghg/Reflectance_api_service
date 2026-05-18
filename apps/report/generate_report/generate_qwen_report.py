#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""报告生成组件。"""
import json
import os
from datetime import datetime
from typing import Dict, Any, List

import requests
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

try:
    from docx import Document
except ImportError as exc:
    raise ImportError(
        "未安装 python-docx，请先执行: pip install python-docx"
    ) from exc


class QwenReportComponent:
    """负责提示词构建、qwen 调用与 Word 生成。"""

    def __init__(
        self,
        api_key: str = "sk-181711d0446247e8bb89ee38aedabaf7",
        model: str = "qwen-plus",
        api_url: str = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
    ):
        self.api_key = api_key
        self.model = model
        self.api_url = api_url

    @staticmethod
    def load_json(path: str) -> Dict[str, Any]:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)

    def build_prompt(self, weather_payload: Dict[str, Any], risk_summary: Dict[str, Any]) -> str:
        all_weather_data: List[Dict[str, Any]] = weather_payload.get("data", [])
        timestamp = risk_summary.get("timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

        prompt = f"""天津市天气预警报告生成任务

任务说明：
你是一位专业的气象分析师，请根据提供的天津市各区域天气数据，生成一份专业的天气预警报告。

数据采集时间：
{timestamp}

整体风险评估：
- 监测区域总数: {risk_summary.get('total_districts', 0)}个
- 存在风险的区域: {len(risk_summary.get('districts_with_warnings', []))}个
- 整体风险等级: {risk_summary.get('overall_risk_level', '未知')}

各区域详细天气数据：
"""

        for data in all_weather_data:
            district_name = data.get("district_name", "未知")
            prompt += f"\n【{district_name}】\n"

            nowinfo = data.get("nowinfo", {})
            if nowinfo:
                prompt += (
                    f"- 当前天气: 温度{nowinfo.get('temperature', 'N/A')}℃, "
                    f"湿度{nowinfo.get('humidity', 'N/A')}%, "
                    f"气压{nowinfo.get('pressure', 'N/A')}hPa, "
                    f"风向{nowinfo.get('windDirection', 'N/A')}, "
                    f"风速{nowinfo.get('windSpeed', 'N/A')}m/s\n"
                )

            prompt += "- 未来3天预报:\n"
            for day_idx in range(1, 4):
                if day_idx == 1:
                    day_data = data
                    date_str = day_data.get("nowdate", "今天")
                else:
                    day_data = data.get(f"weatherday{day_idx}", {})
                    date_str = day_data.get("date", f"第{day_idx}天")

                if not day_data:
                    continue

                weather_day = day_data.get("weather1", "N/A")
                weather_night = day_data.get("weather2", "N/A")
                temp_high = day_data.get("wd1", "N/A")
                temp_low = day_data.get("wd2", "N/A")
                wind_dir = day_data.get("winddirection1", "N/A")
                wind_level = day_data.get("windleve1", "N/A")
                day_label = "今天" if day_idx == 1 else f"第{day_idx}天({date_str})"
                prompt += (
                    f"  - {day_label}: {weather_day}转{weather_night}, "
                    f"温度{temp_low}~{temp_high}℃, {wind_dir}{wind_level}\n"
                )

            prompt += "- 重要时段预报:\n"
            for hour_idx in range(1, 4):
                hour_data = data.get(f"hour{hour_idx}", [])
                if not hour_data:
                    continue
                day_label = "今天" if hour_idx == 1 else f"第{hour_idx}天"
                significant_hours = []
                for hour in hour_data[:6]:
                    rain = hour.get("降水", "")
                    if rain and rain != "无降水":
                        time_str = hour.get("时间", "")
                        weather = hour.get("天气", "")
                        temp = hour.get("气温", "")
                        significant_hours.append(f"{time_str}({weather},{temp},{rain})")
                if significant_hours:
                    prompt += f"  - {day_label}: {'; '.join(significant_hours)}\n"

        if risk_summary.get("districts_with_warnings"):
            prompt += "\n风险预警汇总：\n"
            for district_warning in risk_summary["districts_with_warnings"]:
                prompt += (
                    f"\n【{district_warning.get('district', '未知')}】"
                    f"({district_warning.get('risk_count', 0)}条预警)\n"
                )
                for warning in district_warning.get("warnings", []):
                    prompt += f"- {warning}\n"

        prompt += """

报告输出要求：
请输出结构化报告，包含以下章节：
1. 报告标题（含日期与风险等级）
2. 概要总结（200字以内）
3. 重点预警区域（3-5个）
4. 分时段风险分析
5. 防御建议（市民、行业、农业、交通）
6. 室外作业建议（独立章节，结合本轮天气数据具体分析，须包含但不限于：）
   - 是否适宜开展高空、露天、野外作业及建议停工时段
   - 建筑施工、市政养护、电力与通信巡检、港口与物流装卸、农林牧渔户外作业等分类建议
   - 强降雨、大风、高温、低能见度对应的安全防护与应急措施
   - 每条建议须可执行，与上文风险分析时段相呼应
7. 后续关注要点

要求：
- 语言专业、清晰、可执行
- 所有结论与数据保持一致
- 室外作业建议不得空泛，须引用上文具体天气与风险信息
- 输出直接可用于政府或企业简报
- 只输出纯文本结构，不要使用 Markdown 语法
- 不要出现 #、*、---、|、表格分隔线
"""
        return prompt

    def call_qwen(self, prompt: str) -> str:
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": self.model,
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "你是专业气象预警分析师。"
                        "输出必须是干净的中文纯文本报告，不要 Markdown、不要表格符号、不要分隔线。"
                        "报告必须包含独立的「室外作业建议」章节，内容具体、可执行。"
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.5,
        }
        response = requests.post(self.api_url, headers=headers, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        choices = data.get("choices", [])
        if not choices:
            raise ValueError(f"模型返回异常: {json.dumps(data, ensure_ascii=False)}")
        content = choices[0].get("message", {}).get("content", "")
        if not content:
            raise ValueError("模型返回内容为空")
        return content

    @staticmethod
    def _set_run_font(run, font_name: str, size: int, bold: bool = False) -> None:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size)
        run.bold = bold

    @staticmethod
    def _clean_line(raw_line: str) -> str:
        line = raw_line.strip()
        if not line:
            return ""
        if set(line) <= {"-", "—", "|", " "}:
            return ""
        line = line.replace("**", "").replace("###", "").replace("##", "").replace("#", "").strip()
        if "|" in line:
            columns = [col.strip() for col in line.split("|") if col.strip()]
            line = "；".join(columns)
        return line.strip()

    @staticmethod
    def text_to_docx(report_text: str, output_path: str) -> None:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "宋体"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal_style.font.size = Pt(12)

        title = doc.add_paragraph()
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("天津市天气预警报告")
        QwenReportComponent._set_run_font(title_run, "黑体", 22, bold=True)

        meta = doc.add_paragraph()
        meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta_run = meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        QwenReportComponent._set_run_font(meta_run, "宋体", 11)

        subtitle = doc.add_paragraph()
        subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subtitle_run = subtitle.add_run("发布级别：气象预警分析简报")
        QwenReportComponent._set_run_font(subtitle_run, "宋体", 11)

        section_headers = ("一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、")

        for raw_line in report_text.splitlines():
            line = QwenReportComponent._clean_line(raw_line)
            if not line:
                continue

            if line.startswith(section_headers) or line.startswith("【"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = p.add_run(line)
                QwenReportComponent._set_run_font(run, "黑体", 14, bold=True)
            elif line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(line[2:])
                QwenReportComponent._set_run_font(run, "宋体", 12)
            elif line[:2].isdigit() and line[1:3] in [". ", "、"]:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(line)
                QwenReportComponent._set_run_font(run, "黑体", 12, bold=True)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.84)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(line)
                QwenReportComponent._set_run_font(run, "宋体", 12)
        doc.save(output_path)

    def run(
        self,
        weather_path: str = "tianjin_weather_data.json",
        risk_path: str = "tianjin_risk_analysis.json",
        report_docx_path: str = "qwen_weather_report.docx",
    ) -> str:
        if not os.path.exists(weather_path) or not os.path.exists(risk_path):
            raise FileNotFoundError("缺少输入文件，请先采集天气并生成 JSON 数据")

        weather_payload = self.load_json(weather_path)
        risk_summary = self.load_json(risk_path)

        print("【步骤1】构建提示词并调用 qwen-plus")
        prompt = self.build_prompt(weather_payload, risk_summary)
        report_text = self.call_qwen(prompt)

        print("\n【步骤2】生成 Word")
        self.text_to_docx(report_text, report_docx_path)
        print(f"Word 报告已生成: {report_docx_path}")
        return report_docx_path
